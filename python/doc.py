from docx import Document

# Create a new Word document
doc = Document()

# Title
doc.add_heading('Software Testing Report', level=1)

# Section for Moodle
doc.add_heading('Report for Moodle:', level=2)

# Shortest, Longest, and Average Execution Time
doc.add_heading('Shortest, Longest, and Average Execution Time:', level=3)
table = doc.add_table(rows=4, cols=4)
table.style = 'Table Grid'
table_headers = ['Samples', 'Shortest Time', 'Longest Time', 'Average Time']
table_data = [
    ['50', '33 ms', '717 ms', '278 ms'],
    ['100', '34 ms', '680 ms', '273 ms']
]
for idx, header in enumerate(table_headers):
    table.cell(0, idx).text = header
for row_idx, row in enumerate(table_data, start=1):
    for col_idx, cell in enumerate(row):
        table.cell(row_idx, col_idx).text = cell

doc.add_paragraph()  # Add a blank line

# Comparisons Table
doc.add_heading('Comparisons of Average Throughput, Average Data Reception, and Average Data Transmission:', level=3)
table = doc.add_table(rows=4, cols=4)
table.style = 'Table Grid'
table_headers = ['Samples', 'Average Throughput (req/sec)', 'Average Data Reception (KB)', 'Average Data Transmission (KB)']
table_data = [
    ['50', '2.0', '3.5', '1.335'],
    ['100', '4.0', '3.4775', '1.32']
]
for idx, header in enumerate(table_headers):
    table.cell(0, idx).text = header
for row_idx, row in enumerate(table_data, start=1):
    for col_idx, cell in enumerate(row):
        table.cell(row_idx, col_idx).text = cell

doc.add_paragraph()  # Add a blank line

# Section for Home
doc.add_heading('Report for Home:', level=2)

# Shortest, Longest, and Average Execution Time
doc.add_heading('Shortest, Longest, and Average Execution Time:', level=3)
table = doc.add_table(rows=4, cols=4)
table.style = 'Table Grid'
table_data = [
    ['50', '35 ms', '1033 ms', '275 ms'],
    ['100', '38 ms', '1522 ms', '265 ms']
]
for idx, header in enumerate(table_headers):
    table.cell(0, idx).text = header
for row_idx, row in enumerate(table_data, start=1):
    for col_idx, cell in enumerate(row):
        table.cell(row_idx, col_idx).text = cell

doc.add_paragraph()  # Add a blank line

# Comparisons Table
doc.add_heading('Comparisons of Average Throughput, Average Data Reception, and Average Data Transmission:', level=3)
table = doc.add_table(rows=4, cols=4)
table.style = 'Table Grid'
table_data = [
    ['50', '3.0', '37.57', '0.83'],
    ['100', '6.0', '37.28', '0.825']
]
for idx, header in enumerate(table_headers):
    table.cell(0, idx).text = header
for row_idx, row in enumerate(table_data, start=1):
    for col_idx, cell in enumerate(row):
        table.cell(row_idx, col_idx).text = cell

doc.add_paragraph()  # Add a blank line

# Section for Notice
doc.add_heading('Report for Notice:', level=2)

# Shortest, Longest, and Average Execution Time
doc.add_heading('Shortest, Longest, and Average Execution Time:', level=3)
table = doc.add_table(rows=4, cols=4)
table.style = 'Table Grid'
table_data = [
    ['50', '43 ms', '451 ms', '201 ms'],
    ['100', '42 ms', '480 ms', '196 ms']
]
for idx, header in enumerate(table_headers):
    table.cell(0, idx).text = header
for row_idx, row in enumerate(table_data, start=1):
    for col_idx, cell in enumerate(row):
        table.cell(row_idx, col_idx).text = cell

doc.add_paragraph()  # Add a blank line

# Comparisons Table
doc.add_heading('Comparisons of Average Throughput, Average Data Reception, and Average Data Transmission:', level=3)
table = doc.add_table(rows=4, cols=4)
table.style = 'Table Grid'
table_data = [
    ['50', '1.0', '2.19', '1.65'],
    ['100', '2.0', '2.165', '1.635']
]
for idx, header in enumerate(table_headers):
    table.cell(0, idx).text = header
for row_idx, row in enumerate(table_data, start=1):
    for col_idx, cell in enumerate(row):
        table.cell(row_idx, col_idx).text = cell

# Save the document
output_path = "Revised_Software_Testing_Report.docx"
doc.save(output_path)
print(f"File saved at: {output_path}")